"""
Document Generator for ClarityOS
Creates addressible.docx files with mentor context information
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import json


def create_addressible_docx(
    user_summary: str,
    category: str,
    insights: list,
    metrics: dict,
    questions_for_mentor: list,
    filename: str = None
) -> tuple[str, str]:
    """
    Generate an addressible.docx file with all context for the mentor.
    
    Returns:
        tuple: (file_path, document_content_as_text)
    """
    doc = Document()
    
    # Title
    title = doc.add_heading('🎯 Mentor Context Pack', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    doc.add_paragraph(f"Category: {category}")
    doc.add_paragraph("─" * 50)
    
    # Problem Summary
    doc.add_heading('Problem Summary', level=1)
    doc.add_paragraph(user_summary)
    
    # Key Insights
    doc.add_heading('Key Insights', level=1)
    for i, insight in enumerate(insights, 1):
        doc.add_paragraph(f"{i}. {insight}", style='List Number')
    
    # Metrics (if provided)
    if metrics:
        doc.add_heading('Key Metrics', level=1)
        for key, value in metrics.items():
            doc.add_paragraph(f"• {key}: {value}")
    
    # Questions for Mentor
    doc.add_heading('Questions for Mentor to Address', level=1)
    for q in questions_for_mentor:
        doc.add_paragraph(f"• {q}")
    
    # Footer
    doc.add_paragraph("─" * 50)
    footer = doc.add_paragraph("Powered by ClarityOS | ExpertBells")
    footer.runs[0].font.size = Pt(10)
    
    # Create output directory
    output_dir = os.path.join(os.path.dirname(__file__), "generated_documents")
    os.makedirs(output_dir, exist_ok=True)
    
    # Generate filename
    if not filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"addressible_{timestamp}.docx"
    
    file_path = os.path.join(output_dir, filename)
    doc.save(file_path)
    
    # Create text version for preview
    text_content = f"""
📋 MENTOR CONTEXT PACK
━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Category: {category}
Generated: {datetime.now().strftime('%B %d, %Y')}

📝 PROBLEM SUMMARY:
{user_summary}

💡 KEY INSIGHTS:
{chr(10).join([f'  {i+1}. {ins}' for i, ins in enumerate(insights)])}

📊 METRICS:
{chr(10).join([f'  • {k}: {v}' for k, v in metrics.items()]) if metrics else '  (No metrics provided)'}

❓ QUESTIONS FOR MENTOR:
{chr(10).join([f'  • {q}' for q in questions_for_mentor])}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""
    
    return file_path, text_content


def extract_document_data(conversation_history: list, ai_summary: dict) -> dict:
    """
    Extract structured data from conversation for document generation.
    """
    # Combine all user messages
    user_messages = [m["content"] for m in conversation_history if m.get("role") == "user"]
    
    return {
        "user_summary": ai_summary.get("problem_summary", " ".join(user_messages)),
        "category": ai_summary.get("category", "General"),
        "insights": ai_summary.get("insights", [
            "User is seeking expert guidance",
            "Problem requires specialized mentorship"
        ]),
        "metrics": ai_summary.get("metrics", {}),
        "questions_for_mentor": ai_summary.get("questions_for_mentor", [
            "What specific challenges are you facing?",
            "What outcomes do you expect from this session?",
            "What constraints do you have (time, budget, resources)?"
        ])
    }
